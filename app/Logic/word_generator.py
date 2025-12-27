from docx import Document
from datetime import datetime

MESES_ES = {
    1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril",
    5: "Mayo", 6: "Junio", 7: "Julio", 8: "Agosto",
    9: "Septiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"
}

def generar_documento_word(
    ruta_documento_base: str,
    ruta_salida: str,
    datos_cliente: dict
):
    """
    Genera un documento Word usando como base un documento existente
    y reemplazando marcadores con los datos del cliente.
    """

    fecha = datetime.now()

    datos = datos_cliente.copy()
    datos["DIA"] = str(fecha.day)
    datos["MES"] = MESES_ES[fecha.month]
    datos["ANIO"] = str(fecha.year)

    doc = Document(ruta_documento_base)

    def reemplazar_texto(parrafo):
        texto = "".join(run.text for run in parrafo.runs)

        for clave, valor in datos.items():
            marcador = f"{{{{{clave}}}}}"
            texto = texto.replace(marcador, str(valor))

        if parrafo.runs:
            parrafo.runs[0].text = texto
            for run in parrafo.runs[1:]:
                run.text = ""

    # Párrafos normales
    for p in doc.paragraphs:
        reemplazar_texto(p)

    # Tablas
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    reemplazar_texto(p)

    doc.save(ruta_salida)
