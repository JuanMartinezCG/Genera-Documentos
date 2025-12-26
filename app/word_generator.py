from docx import Document
from datetime import datetime
import os

def generar_documento_word(datos_cliente: dict, nombre_salida: str):
    """Genera un documento de Word a partir de una plantilla y datos del cliente."""

    MESES_ES = {
    1: "Enero",
    2: "Febrero",
    3: "Marzo",
    4: "Abril",
    5: "Mayo",
    6: "Junio",
    7: "Julio",
    8: "Agosto",
    9: "Septiembre",
    10: "Octubre",
    11: "Noviembre",
    12: "Diciembre"
    }

    fecha_actual = datetime.now()

    datos_cliente = datos_cliente.copy()
    datos_cliente["DIA"] = str(fecha_actual.day)
    datos_cliente["MES"] = MESES_ES[fecha_actual.month]
    datos_cliente["ANIO"] = str(fecha_actual.year)

    base_dir = os.path.dirname(os.path.abspath(__file__)) # Directorio base del script
    ruta_plantilla = os.path.join(base_dir, "plantilla.docx") # Ruta de la plantilla de Word

    doc = Document(ruta_plantilla) # Cargar la plantilla de Word

    def reemplazar_texto_en_parrafo(parrafo, datos):
        for clave, valor in datos.items(): # Iterar sobre cada par clave-valor en los datos del cliente
            marcador = f"{{{{{clave}}}}}" # Formatear el marcador de posición
            if marcador in parrafo.text: # Si el marcador está en el texto del párrafo
                parrafo.text = parrafo.text.replace( # lo remplazamos
                    marcador, 
                    str(valor)
                )

    # Texto normal
    for parrafo in doc.paragraphs: # Iterar sobre cada párrafo en el documento
        reemplazar_texto_en_parrafo(parrafo, datos_cliente)

    # Tablas
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    reemplazar_texto_en_parrafo(parrafo, datos_cliente)

    ruta_salida = os.path.join(base_dir, nombre_salida) # Ruta de salida del documento generado
    doc.save(ruta_salida) # Guardar el documento generado 